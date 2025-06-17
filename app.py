import streamlit as st
import os
import pandas as pd
import numpy as np
from sklearn.ensemble import IsolationForest
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
from groq import Groq
from dotenv import load_dotenv
from docx import Document

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

# Set Streamlit page config
st.set_page_config(page_title="𝐌𝐚𝐫𝐯𝐞𝐥.𝐚𝐢 𝐆𝐞𝐧 𝐀𝐈-𝐩𝐨𝐰𝐞𝐫𝐞𝐝 𝐀𝐧𝐨𝐦𝐚𝐥𝐲 𝐏𝐫𝐞𝐝𝐢𝐜𝐭𝐨𝐫", layout="wide")

# Sidebar for file upload
with st.sidebar:
    st.title("📊 Data Ingestion")
    uploaded_file = st.file_uploader("📂 Load your CSV/XLSX returns data", type=["csv", "xlsx"])
    st.markdown("---")

# Main App Title
st.title("🧠𝐌𝐚𝐫𝐯𝐞𝐥.𝐚𝐢 𝐆𝐞𝐧 𝐀𝐈-𝐩𝐨𝐰𝐞𝐫𝐞𝐝 𝐀𝐧𝐨𝐦𝐚𝐥𝐲 𝐏𝐫𝐞𝐝𝐢𝐜𝐭𝐨𝐫")

if uploaded_file is not None:
    # Load data
    if uploaded_file.name.endswith('.csv'):
        data = pd.read_csv(uploaded_file)
    else:
        data = pd.read_excel(uploaded_file)

    st.subheader("📄 Dataset Preview")
    st.dataframe(data.head())
    st.write(f"Total Records: {data.shape[0]} | Total Columns: {data.shape[1]}")

    data.fillna(0, inplace=True)

    # Auto map column names
    col_map = {
        'User_ID': 'customer_id',
        'Order_ID': 'order_id',
        'Return_Reason': 'return_reason',
        'Product_Category': 'category',
        'Product_ID': 'product_id',
        'Product_Price': 'order_amount',
        'Order_Date': 'order_date',
        'Return_Date': 'return_date'
    }
    data.rename(columns=col_map, inplace=True)

    # Check required columns after mapping
    expected_cols = ['customer_id', 'order_id', 'return_reason', 'category', 'product_id',
                     'order_amount', 'return_date', 'order_date']
    missing_cols = [col for col in expected_cols if col not in data.columns]
    if missing_cols:
        st.warning(f"❌ Missing columns after mapping: {missing_cols}. The model may not work as intended.")
        st.stop()

    # Feature Engineering
    data['order_date'] = pd.to_datetime(data['order_date'], errors='coerce')
    data['return_date'] = pd.to_datetime(data['return_date'], errors='coerce')
    data['days_to_return'] = (data['return_date'] - data['order_date']).dt.days

    customer_features = data.groupby('customer_id').agg({
        'order_id': 'count',
        'order_amount': 'sum',
        'days_to_return': 'mean'
    }).rename(columns={'order_id': 'total_returns', 'order_amount': 'total_spent'})
    customer_features.fillna(0, inplace=True)

    # Anomaly Detection
    st.subheader("🚀 Running Anomaly Detection")
    model = IsolationForest(contamination=0.05, random_state=42)
    customer_features['anomaly_score'] = model.fit_predict(customer_features)
    customer_features['return_risk'] = np.where(customer_features['anomaly_score'] == -1, 'High Risk', 'Normal')
    st.success("✅ Anomaly Detection Completed")

    risky_customers = customer_features[customer_features['return_risk'] == 'High Risk']

    # 📊 Key Insights Dashboard
    st.subheader("📊 Key Insights Dashboard")

    total_customers = customer_features.shape[0]
    total_risk_customers = risky_customers.shape[0]
    top_reasons = data['return_reason'].value_counts().head(3).to_dict()
    top_categories = data['category'].value_counts().head(3).to_dict()

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Customers", total_customers)
    col2.metric("🚩 High-Risk Customers", total_risk_customers, help="Customers showing abnormal return behavior")
    col3.metric("Top Return Reason", list(top_reasons.keys())[0], top_reasons[list(top_reasons.keys())[0]])
    col4.metric("Top Return Category", list(top_categories.keys())[0], top_categories[list(top_categories.keys())[0]])

    st.markdown("""
    <div style="background-color:#f1f1f1;padding:10px;border-radius:10px">
    <b>🔎 Explanation:</b> <br>
    High-Risk Customers are flagged by anomaly detection models as potential fraud or unusual return behavior (example: too many returns, high order value returns, late returns, etc.)
    </div>
    """, unsafe_allow_html=True)

    # 🚩 High-Risk Customers Detailed View
    st.subheader("🚩 High-Risk Customers Details")
    st.dataframe(risky_customers.reset_index())
    
    # Download button
    csv = risky_customers.reset_index().to_csv(index=False).encode('utf-8')
    st.download_button(
        label="📥 Download High-Risk Customers (CSV)",
        data=csv,
        file_name='high_risk_customers.csv',
        mime='text/csv'
    )

    # Visual Insights
    st.subheader("📈 Visual Insights")
    col1, col2 = st.columns(2)

    with col1:
        reason_counts = data['return_reason'].value_counts()
        fig1 = px.pie(names=reason_counts.index, values=reason_counts.values, title="Return Reasons Distribution")
        st.plotly_chart(fig1)

    with col2:
        category_counts = data['category'].value_counts()
        fig2 = px.bar(x=category_counts.index, y=category_counts.values, labels={'x':'Category', 'y':'Returns Count'},
                      title="Returns by Product Category")
        st.plotly_chart(fig2)

    return_trends = data.groupby(data['return_date'].dt.to_period('M')).size()
    return_trends.index = return_trends.index.astype(str)
    fig3 = px.line(x=return_trends.index, y=return_trends.values, labels={'x':'Month', 'y':'Returns'},
                   title="Return Trends Over Time")
    st.plotly_chart(fig3, use_container_width=True)

    st.subheader("🔬 Feature Correlation Heatmap")
    plt.figure(figsize=(8, 5))
    sns.heatmap(customer_features[['total_returns','total_spent','days_to_return']].corr(), annot=True, cmap='coolwarm')
    st.pyplot(plt)

    # Prepare insights for LLM
    insights = f"""
Dataset Insights Summary:
- Total Customers: {total_customers}
- High-Risk Customers: {total_risk_customers}
- Top Return Reasons: {top_reasons}
- Top Return Categories: {top_categories}
"""

    st.subheader("💡 AI-Generated Executive Summary")

    if st.button("📝 Generate LLM Summary & Download Word Report"):
        with st.spinner("Generating insights using Marvel.ai GenAI..."):

            prompt = f"""
You are an expert fraud analyst for an e-commerce platform. Based on the following return data insights, generate a professional executive summary highlighting possible fraud patterns, customer behavior anomalies, and product-level concerns. Keep it clear, insightful, and aligned to business language:

{insights}
            """

            completion = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[{"role": "user", "content": prompt}],
                temperature=1,
                max_completion_tokens=1024,
                top_p=1,
                stream=False
            )

            full_summary = completion.choices[0].message.content
            st.write(full_summary)

            # Word Doc Generation
            doc = Document()
            doc.add_heading('Marvel.ai – Executive Summary Report', 0)
            doc.add_paragraph(full_summary)

            doc_path = "executive_summary.docx"
            doc.save(doc_path)

            with open(doc_path, "rb") as file:
                st.download_button("📄 Download Word Report", data=file, file_name=doc_path)

else:
    st.info("✨ Upload your data to let Marvel.ai get started")
